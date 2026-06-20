import re
from io import BytesIO
from typing import List, Dict, Any, Optional
from datetime import datetime
import spacy
from pypdf import PdfReader
from docx import Document

class ResumeParserService:
    def __init__(self):
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except Exception:
            self.nlp = None

        # Regex patterns
        self.email_pattern = re.compile(r'[\w\.-]+@[\w\.-]+\.\w+')
        self.phone_pattern = re.compile(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
        
        # Comprehensive tech and soft skills dictionary
        self.SKILLS_DB = {
            # Languages
            "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "php", "go", "golang", "rust", "swift", "kotlin", "scala", "r", "sql", "html", "css",
            # Frameworks & Libraries
            "react", "angular", "vue", "next.js", "nextjs", "nuxt", "svelte", "django", "fastapi", "flask", "spring boot", "spring", "express", "expressjs", "nest.js", "nestjs", "laravel", "rails", "asp.net",
            # Data Science & ML
            "pytorch", "tensorflow", "keras", "scikit-learn", "sklearn", "numpy", "pandas", "scipy", "nltk", "spacy", "opencv", "huggingface", "transformers", "llm", "langchain", "llama",
            # Databases
            "postgresql", "postgres", "mysql", "mongodb", "redis", "sqlite", "cassandra", "dynamodb", "neo4j", "elasticsearch", "oracle",
            # DevOps & Cloud
            "docker", "kubernetes", "k8s", "aws", "amazon web services", "azure", "gcp", "google cloud", "terraform", "ansible", "jenkins", "github actions", "gitlab ci", "ci/cd", "vagrant",
            # Concepts / Methods
            "rest api", "graphql", "grpc", "microservices", "agile", "scrum", "kanban", "git", "linux", "bash", "system design", "oop", "mvc",
            # Soft Skills / Business
            "project management", "product management", "leadership", "communication", "teamwork", "problem solving", "time management", "analytical skills"
        }

        # Degrees mapping
        self.degree_keywords = {
            "bachelor": ["bachelor", "b.s", "bs", "b.tech", "btech", "b.e", "be", "b.a", "ba", "bsc", "bca"],
            "master": ["master", "m.s", "ms", "m.tech", "mtech", "m.e", "me", "m.a", "ma", "msc", "mca", "mba"],
            "phd": ["phd", "ph.d", "doctorate", "doctor of philosophy"]
        }

    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        text = ""
        try:
            reader = PdfReader(BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"Error parsing PDF: {e}")
        return text

    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        text = ""
        try:
            doc = Document(BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
        return text

    def parse_resume(self, text: str, filename: str) -> Dict[str, Any]:
        """
        Parses resume text and extracts contact details, skills, education, experience, projects, and certifications.
        """
        # Cleanup text slightly
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        clean_text = " ".join(lines)

        # 1. Extract Email
        email = None
        emails = self.email_pattern.findall(clean_text)
        if emails:
            email = emails[0]

        # 2. Extract Phone
        phone = None
        phones = self.phone_pattern.findall(clean_text)
        if phones:
            phone = phones[0]

        # 3. Extract Name
        name = self._extract_name(text, lines)

        # 4. Extract Skills
        skills = self._extract_skills(clean_text)

        # 5. Extract Education
        education = self._extract_education(text)

        # 6. Extract Experience
        experience = self._extract_experience(text)

        # 7. Extract Projects
        projects = self._extract_projects(text)

        # 8. Extract Certifications
        certifications = self._extract_certifications(text)

        return {
            "name": name or filename.split(".")[0],
            "email": email,
            "phone": phone,
            "skills": skills,
            "education": education,
            "experience": experience,
            "projects": projects,
            "certifications": certifications
        }

    def _extract_name(self, raw_text: str, lines: List[str]) -> Optional[str]:
        if not self.nlp or not lines:
            return None
        
        # Usually, the name is in the first 3 lines
        header_text = "\n".join(lines[:3])
        doc = self.nlp(header_text)
        
        # Look for PERSON entities
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                # Ensure the name doesn't contain emails or links
                val = ent.text.strip()
                if "@" not in val and "/" not in val and len(val.split()) >= 2 and len(val) < 40:
                    return val
        
        # Fallback to the first line if it looks like a name
        first_line = lines[0]
        if len(first_line.split()) >= 2 and len(first_line) < 30 and not any(char.isdigit() for char in first_line) and "@" not in first_line:
            return first_line

        return None

    def _extract_skills(self, text: str) -> List[str]:
        # Perform exact keyword matching (case-insensitive)
        text_lower = text.lower()
        extracted_skills = []
        for skill in self.SKILLS_DB:
            # Match using word boundaries to avoid sub-word matching (e.g. 'go' matching 'google')
            # For special skills like c++, c#, .net, next.js, we adjust boundaries
            pattern = rf"\b{re.escape(skill)}\b"
            if "+" in skill or "#" in skill or "." in skill:
                pattern = rf"(?:^|\s|\W){re.escape(skill)}(?:$|\s|\W)"
            
            if re.search(pattern, text_lower):
                extracted_skills.append(skill)
        
        # Deduplicate and sort
        return sorted(list(set(extracted_skills)))

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        education_list = []
        lines = text.split("\n")
        
        # Look for sections containing education keyword
        edu_keywords = ["education", "academic", "qualification", "university", "college", "schooling"]
        is_edu_section = False
        edu_lines = []
        
        for line in lines:
            line_lower = line.lower().strip()
            # If we hit a new major header after education, stop
            if is_edu_section and any(h in line_lower for h in ["experience", "work", "history", "projects", "skills", "certifications", "interests"]):
                is_edu_section = False
            
            if any(k in line_lower for k in edu_keywords) and len(line_lower) < 25:
                is_edu_section = True
                continue
                
            if is_edu_section:
                edu_lines.append(line)
        
        # If no explicit education section, search the entire text for degrees
        lines_to_search = edu_lines if edu_lines else lines
        
        for line in lines_to_search:
            line_lower = line.lower()
            found_degree = None
            degree_level = None
            
            for level, aliases in self.degree_keywords.items():
                for alias in aliases:
                    if re.search(rf"\b{re.escape(alias)}\b", line_lower):
                        found_degree = line.strip()
                        degree_level = level
                        break
                if found_degree:
                    break
            
            if found_degree:
                # Try to extract year in the line
                year_match = re.search(r"\b(19|20)\d{2}\b", line)
                year = year_match.group(0) if year_match else None
                education_list.append({
                    "degree": found_degree,
                    "level": degree_level,
                    "year": year
                })

        return education_list

    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        experience_list = []
        lines = text.split("\n")
        
        # Identify experience section
        exp_keywords = ["experience", "employment", "work history", "professional background", "career history", "work experience"]
        is_exp_section = False
        exp_lines = []
        
        for line in lines:
            line_lower = line.lower().strip()
            if is_exp_section and any(h in line_lower for h in ["education", "academic", "projects", "skills", "certifications", "interests"]):
                is_exp_section = False
            
            if any(k in line_lower for k in exp_keywords) and len(line_lower) < 25:
                is_exp_section = True
                continue
                
            if is_exp_section:
                exp_lines.append(line.strip())
                
        # Simple extraction: look for job title patterns or company patterns
        # Standard format often lists Job Title, Company Name, Dates (e.g. "Software Engineer - Google (2018 - 2021)")
        # We can extract non-empty lines in the section as experience bullet points/records
        current_exp = None
        for line in exp_lines:
            if not line:
                continue
            
            # If line contains year range, it's likely a job entry header
            # Examples: 2018-2020, 2021 - Present, Jan 2020 to Aug 2022
            date_range_match = re.search(r"\b(?:20|19)\d{2}\s*[-–to\s]+\s*(?:present|(?:20|19)\d{2}|current)\b", line, re.IGNORECASE)
            
            if date_range_match:
                if current_exp:
                    experience_list.append(current_exp)
                current_exp = {
                    "role_details": line,
                    "duration": date_range_match.group(0),
                    "responsibilities": []
                }
            elif current_exp:
                # Append subsequent lines as responsibilities/details
                if len(line) > 10:
                    current_exp["responsibilities"].append(line)
        
        if current_exp:
            experience_list.append(current_exp)
            
        # Fallback if no formatted section found: look for years of experience statements
        # e.g., "5 years of experience in python"
        return experience_list

    def _extract_projects(self, text: str) -> List[Dict[str, Any]]:
        project_list = []
        lines = text.split("\n")
        
        proj_keywords = ["projects", "personal projects", "academic projects", "key projects"]
        is_proj_section = False
        proj_lines = []
        
        for line in lines:
            line_lower = line.lower().strip()
            if is_proj_section and any(h in line_lower for h in ["education", "experience", "work", "skills", "certifications", "interests"]):
                is_proj_section = False
            
            if any(k in line_lower for k in proj_keywords) and len(line_lower) < 25:
                is_proj_section = True
                continue
                
            if is_proj_section:
                proj_lines.append(line.strip())

        # Collect project titles and descriptions
        current_project = None
        for line in proj_lines:
            if not line:
                continue
            # If line is short and bold/capitalized (likely title), or starting with bullet points
            if len(line) < 50 and not line.startswith("-") and not line.startswith("*"):
                if current_project:
                    project_list.append(current_project)
                current_project = {
                    "title": line,
                    "description": ""
                }
            elif current_project:
                current_project["description"] += line + " "
                
        if current_project:
            project_list.append(current_project)
            
        return project_list

    def _extract_certifications(self, text: str) -> List[str]:
        certifications = []
        lines = text.split("\n")
        
        cert_keywords = ["certifications", "licenses", "courses", "certificates"]
        is_cert_section = False
        
        for line in lines:
            line_lower = line.lower().strip()
            if is_cert_section and any(h in line_lower for h in ["education", "experience", "work", "projects", "skills", "interests"]):
                is_cert_section = False
            
            if any(k in line_lower for k in cert_keywords) and len(line_lower) < 25:
                is_cert_section = True
                continue
                
            if is_cert_section and len(line.strip()) > 5:
                certifications.append(line.strip())
                
        # If no explicit section, look for common certificates anywhere in the text
        if not certifications:
            common_certs = [
                "aws certified", "google cloud certified", "azure certified", "pmp", "scrum master", "certified scrummaster",
                "cisco", "ccna", "ccnp", "comptia", "itil", "oracle certified"
            ]
            text_lower = text.lower()
            for cert in common_certs:
                if cert in text_lower:
                    # Find matching line
                    for line in lines:
                        if cert in line.lower() and len(line.strip()) < 80:
                            certifications.append(line.strip())
                            break
                            
        return list(set(certifications))

    @staticmethod
    def calculate_years_of_experience(experience_list: List[Dict[str, Any]], raw_text: str) -> float:
        """
        Parses dates to calculate total years of experience, fallback to scanning text for statements like "X years of experience".
        """
        total_years = 0.0
        
        # 1. Parse from experience list date ranges
        for exp in experience_list:
            duration = exp.get("duration", "")
            # Look for years
            years = re.findall(r"\b(19|20)\d{2}\b", duration)
            if len(years) == 2:
                try:
                    diff = int(years[1]) - int(years[0])
                    total_years += max(0, diff)
                except ValueError:
                    pass
            elif len(years) == 1 and ("present" in duration.lower() or "current" in duration.lower()):
                try:
                    diff = datetime.now().year - int(years[0])
                    total_years += max(0, diff)
                except Exception:
                    # Default 1 year if parsing fails
                    total_years += 1.0

        if total_years > 0:
            return total_years

        # 2. Regex fallback: e.g., "5+ years of experience" or "worked for 4 years"
        match = re.search(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:work\s*)?experience", raw_text, re.IGNORECASE)
        if match:
            return float(match.group(1))
            
        return 0.0
